import docx
doc = docx.Document()
doc.add_heading("Header 0", 0)
doc.add_heading("Header 1", 1)
doc.add_heading("Header 2", 2)
doc.add_heading("Header 3", 3)
doc.add_picture("OU-logo.png", width=docx.shared.Cm(4), height=docx.shared.Cm(4))
doc.add_paragraph("Hello World!", "Title")

para_1 = doc.add_paragraph("This is the second paragraph.")
para_1.add_run("Adding run to the second paragraph.")
para_1.runs[0].add_text("Adding text for run.")
para_1.runs[1].add_picture("OU-logo.png", width=docx.shared.Inches(2), height=docx.shared.Cm(4))
print(f'len(para_1.runs) = {len(para_1.runs)}')

para_1.runs[1].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.add_paragraph("This is on the second page!")
doc.save("doc1.docx")
